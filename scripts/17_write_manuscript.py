"""
Generates the full HPAI Bangladesh manuscript as a formatted Word (.docx) file.
Output: /Users/khalilur/Documents/AIWORK/avian-influenza/paper/Bangladesh_HPAI_Manuscript.docx

All numbers in this script are pulled from data/table1_division_summary.csv,
data/table2_covariate_correlation.csv, data/table2b_seasonal_split.csv,
data/table3_negbin_regression.csv, and data/processed/poultry_density_by_division.csv
-- nothing here is invented. References are limited to organizational
sources (WOAH/FAO/WHO/NASA/geoBoundaries) that this pipeline actually used,
plus a small number of Bangladesh HPAI journal citations that MUST be
independently verified before submission (flagged at the end of this file
and in the printed output).
"""

from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path

BASE_DIR = Path(__file__).parent.parent
FIG_DIR = BASE_DIR / "figures"
OUT = BASE_DIR / "paper" / "Bangladesh_HPAI_Manuscript.docx"
OUT.parent.mkdir(parents=True, exist_ok=True)

# ── Document setup ──────────────────────────────────────────────────────────
doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style.font.size = Pt(12)


# ── Helper functions (same pattern as bangladesh-dengue-2018-2025) ─────────

def heading(text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.runs[0] if p.runs else p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.bold = True
    run.font.size = Pt(13) if level == 1 else Pt(12)
    return p


def para(text, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=6):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = Pt(22)
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    run.bold = bold
    run.italic = italic
    return p


def para_indent(text):
    p = para(text)
    p.paragraph_format.first_line_indent = Cm(1.27)
    return p


def mixed_para(*parts):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = Pt(22)
    for text, bold, italic in parts:
        run = p.add_run(text)
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
        run.bold = bold
        run.italic = italic
    return p


def add_figure(path, caption, width=Inches(5.5)):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=width)
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(12)
    run2 = cap.add_run(caption)
    run2.font.name = "Times New Roman"
    run2.font.size = Pt(10)
    run2.italic = True


def add_table_from_data(headers, rows, caption):
    cap = doc.add_paragraph()
    cap.paragraph_format.space_before = Pt(6)
    r = cap.add_run(caption)
    r.font.name = "Times New Roman"
    r.font.size = Pt(10)
    r.bold = True

    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = "Table Grid"

    hdr = tbl.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
            run.font.name = "Times New Roman"
            run.font.size = Pt(10)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for r_idx, row_data in enumerate(rows):
        row = tbl.rows[r_idx + 1]
        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = str(val)
            for run in cell.paragraphs[0].runs:
                run.font.name = "Times New Roman"
                run.font.size = Pt(10)

    doc.add_paragraph()


def page_break():
    doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ══════════════════════════════════════════════════════════════════════════

title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_p.paragraph_format.space_before = Pt(48)
title_p.paragraph_format.space_after = Pt(24)
tr = title_p.add_run(
    "Two Decades of Highly Pathogenic Avian Influenza in Bangladesh: "
    "Spatiotemporal Patterns, Seasonality, and Emerging Wildlife Detections, 2007-2025"
)
tr.font.name = "Times New Roman"
tr.font.size = Pt(14)
tr.font.bold = True

auth_p = doc.add_paragraph()
auth_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
auth_p.paragraph_format.space_after = Pt(6)
ar = auth_p.add_run("Khalilur Rahman Ridoy Khan")
ar.font.name = "Times New Roman"
ar.font.size = Pt(12)

aff_p = doc.add_paragraph()
aff_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
aff_p.paragraph_format.space_after = Pt(6)
afr = aff_p.add_run("Independent Researcher, Dhaka, Bangladesh")
afr.font.name = "Times New Roman"
afr.font.size = Pt(11)
afr.italic = True

corr_p = doc.add_paragraph()
corr_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
corr_p.paragraph_format.space_after = Pt(48)
cr = corr_p.add_run("Correspondence: khalilurrahmanridoykhan@gmail.com")
cr.font.name = "Times New Roman"
cr.font.size = Pt(11)

wc_p = doc.add_paragraph()
wc_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
wcr = wc_p.add_run(
    "Running title: Bangladesh HPAI Spatiotemporal Analysis 2007-2025\n"
    "Manuscript type: Original Research Article (One Health / Veterinary Epidemiology)"
)
wcr.font.name = "Times New Roman"
wcr.font.size = Pt(11)

page_break()

# ══════════════════════════════════════════════════════════════════════════
# ABSTRACT
# ══════════════════════════════════════════════════════════════════════════

heading("Abstract", 1)

mixed_para(("Background: ", True, False),
           ("Highly pathogenic avian influenza (HPAI) has recurred in Bangladesh's "
            "poultry sector since 2007, with periodic detections in wild birds since "
            "2017. Most published analyses of Bangladesh HPAI cover only the initial "
            "2007-2013 epidemic wave using farm-level survey data. A systematic "
            "characterisation using the full official international surveillance "
            "record through 2025 -- including recent wildlife detections -- has not "
            "been available.", False, False))

mixed_para(("Methods: ", True, False),
           ("We built a reproducible pipeline on three linked data products from the "
            "World Organisation for Animal Health (WOAH) World Animal Health "
            "Information System (WAHIS) for Bangladesh: six-month country reports, "
            "an event-level quantitative-data export (division, semester, species, "
            "and serotype resolution), and individual outbreak/event records "
            "(2007-2025). Division-level climate covariates (temperature, relative "
            "humidity, precipitation; NASA POWER) and a static 2020 poultry-density "
            "layer (FAO Gridded Livestock of the World v4) were merged with outbreak "
            "counts. Associations were tested with Pearson correlation and a "
            "negative-binomial regression with division fixed effects.", False, False))

mixed_para(("Results: ", True, False),
           ("Bangladesh reported 564 new HPAI outbreaks across 8 administrative "
            "divisions/localities from 2007 to 2025, with 465,003 susceptible-animal "
            "cases, 2,391,300 birds killed/disposed of under control measures, and "
            "447,872 recorded deaths. Dhaka division carried the highest burden (266 "
            "outbreaks; 267 including a 2025 event in Narayanganj Sadar), followed by "
            "Rajshahi (131) and Chittagong (81); Sylhet recorded the fewest (4). "
            "Outbreaks were strongly seasonal: 525 of 563 division-resolved outbreaks "
            "(93.3%) occurred in the January-June (dry/cool) semester versus 38 in "
            "July-December (monsoon). Temperature, humidity, and precipitation were "
            "each significantly associated with outbreak counts (p<0.05) after "
            "controlling for division in a negative-binomial model (pseudo R2 = "
            "0.136). A finer-resolution reconstruction of the largest tracked event "
            "(2007-2013, 549 cumulative outbreaks across 46 follow-up reports) "
            "revealed two distinct epidemic peaks -- April 2008 (156 new outbreaks in "
            "one report) and March 2011 (67) -- separated by a quiet 2009-2010 "
            "period, which is invisible in semester-aggregated data. Wild-species "
            "detections (House Crow, unidentified Phasianidae) were recorded "
            "2016-2019, and the most recent record (April 2025) was H5N1 in a "
            "captive Serval (wild cat) in Narayanganj, with 2 of 2 susceptible "
            "animals dying.", False, False))

mixed_para(("Conclusions: ", True, False),
           ("Bangladesh's HPAI burden is concentrated in Dhaka division and strongly "
            "concentrated in the dry season, a pattern that persists after "
            "controlling for division-level differences and that is not explained "
            "by poultry density alone. The 2025 captive-Serval detection illustrates "
            "an emerging One Health surveillance need -- confirmed spillover into "
            "non-poultry hosts, distinct from the historical poultry-sector pattern "
            "-- that warrants closer integration of veterinary, wildlife, and public "
            "health surveillance in Bangladesh.", False, False))

mixed_para(("Keywords: ", True, False),
           ("avian influenza; H5N1; Bangladesh; One Health; WAHIS; poultry; "
            "seasonality; zoonotic disease; wildlife spillover", False, True))

page_break()

# ══════════════════════════════════════════════════════════════════════════
# 1. INTRODUCTION
# ══════════════════════════════════════════════════════════════════════════

heading("1. Introduction", 1)

para_indent(
    "Highly pathogenic avian influenza (HPAI), predominantly subtype H5N1, has "
    "caused recurring epizootics in domestic poultry across South and Southeast "
    "Asia since the early 2000s, with periodic spillover into wild birds and, more "
    "recently, mammals [1,2]. Bangladesh, one of the world's most densely "
    "populated countries with a large and predominantly informal live-bird-market "
    "poultry sector, reported its first HPAI outbreaks in early 2007 and has since "
    "experienced repeated epidemic waves [3,4]."
)

para_indent(
    "Most published characterisations of HPAI in Bangladesh focus on the initial "
    "2007-2013 epidemic wave, drawing on farm-level surveillance and risk-factor "
    "studies conducted during that period [3,4]. Since 2017, the World Organisation "
    "for Animal Health (WOAH) has separately tracked HPAI in non-poultry and wild-"
    "bird hosts, and Bangladesh's WAHIS record now extends through 2025, including "
    "a 2025 detection in a captive wild mammal. A systematic, reproducible "
    "characterisation spanning the full 2007-2025 WAHIS record -- integrating "
    "outbreak counts, species involved, and environmental covariates -- has not, "
    "to our knowledge, been previously reported."
)

para_indent(
    "This gap matters for two reasons. First, HPAI H5N1 remains a priority "
    "zoonotic and pandemic-preparedness concern; the same virus family has caused "
    "sporadic severe human infections regionally, and understanding the spatial "
    "and seasonal structure of poultry-sector risk in Bangladesh has direct "
    "relevance for early-warning system design. Second, the global emergence since "
    "2020 of H5N1 clade 2.3.4.4b, with documented spillover into a widening range "
    "of wild and domestic mammals internationally, makes any confirmed non-poultry "
    "detection -- such as the 2025 Bangladesh event described here -- worth "
    "characterising carefully within a One Health framework, even from a single "
    "case, rather than only after a larger cluster emerges."
)

para_indent(
    "This study aimed to: (1) characterise the temporal and spatial distribution "
    "of officially reported HPAI outbreaks in Bangladesh, 2007-2025; (2) quantify "
    "the association between outbreak occurrence and division-level climate and "
    "poultry-density covariates; (3) reconstruct a finer-resolution epidemic curve "
    "for the largest tracked outbreak wave using individual report dates; and (4) "
    "document the species involved in reported events, including the most recent "
    "non-poultry detection."
)

# ══════════════════════════════════════════════════════════════════════════
# 2. METHODS
# ══════════════════════════════════════════════════════════════════════════

heading("2. Methods", 1)

heading("2.1 Study Design and Data Sources", 2)
para_indent(
    "We conducted a retrospective descriptive and ecological analysis of official "
    "HPAI surveillance data for Bangladesh reported to WOAH through the World "
    "Animal Health Information System (WAHIS, https://wahis.woah.org), accessed "
    "August 2026. WAHIS aggregates three linked data products used in this study: "
    "(i) six-month country reports (Analytics > Disease situation), giving a "
    "national presence/absence/no-information status per disease per semester "
    "since 2005; (ii) an event-level quantitative-data export (Analytics > "
    "Quantitative data), resolved to administrative division, semester, animal "
    "category, species, and serotype/subtype, with new-outbreak counts and "
    "susceptible/case/death/culling numbers; and (iii) individual animal-disease "
    "events (Reports > Animal disease events), each comprising one or more dated "
    "follow-up reports and, at the individual-outbreak level, exact coordinates. "
    "We restricted all analyses to the two WAHIS disease listings covering HPAI: "
    "'High pathogenicity avian influenza viruses (Inf. with) (poultry)' and "
    "'Influenza A viruses of high pathogenicity (Inf. with) (non-poultry including "
    "wild birds) (2017-)', the latter introduced when WOAH separated non-poultry "
    "hosts into its own listing in 2017."
)

heading("2.2 Data Extraction and Cleaning", 2)
para_indent(
    "WAHIS's public dashboards render as client-side applications behind bot "
    "protection (Cloudflare on the WAHIS domain; a Flutter web application with no "
    "discoverable public API on the companion FAO EMPRES-i+ platform), so all "
    "exports were retrieved through the standard browser export functions rather "
    "than programmatic scraping. Placeholder values ('-') in the quantitative-data "
    "export denote 'not reported' and were retained as missing (not coerced to "
    "zero) throughout. New-outbreak counts, which WAHIS records on whichever row "
    "of a division/semester/species group carries a value rather than "
    "consistently on an aggregate row, were reconciled by summing within each "
    "group. Bangladesh's current 8-division administrative boundary (geoBoundaries "
    "ADM1 [5]) was dissolved to the 7-division scheme used by WAHIS's division "
    "field, merging Mymensingh (split from Dhaka in 2015) back into Dhaka; the one "
    "outbreak record resolved to the upazila 'Narayanganj Sadar' was assigned to "
    "Dhaka division for mapping and division-level covariate purposes."
)

heading("2.3 Climate and Poultry-Density Covariates", 2)
para_indent(
    "Division-centroid daily temperature, relative humidity, and precipitation "
    "(2007-2025) were obtained from the NASA POWER point API (power.larc.nasa.gov "
    "[6]) and aggregated to WAHIS reporting semesters. Chicken density (head/km2) "
    "was obtained from the FAO Gridded Livestock of the World, version 4 (GLW4), "
    "2020, 10 km resolution [7], and zonal-averaged per division. Because GLW4 "
    "provides a single cross-sectional year, poultry density was used as a static "
    "relative-exposure covariate, not a year-matched population estimate; "
    "Bangladesh's poultry sector is known to have grown substantially over the "
    "2007-2025 study period, so absolute headcounts derived from this layer should "
    "not be read as historical estimates."
)

heading("2.4 Statistical Analysis", 2)
para_indent(
    "Pairwise associations between semester-level outbreak counts and climate/"
    "poultry covariates were assessed with Pearson correlation. A negative-"
    "binomial regression (new outbreaks ~ temperature + humidity + precipitation "
    "+ division fixed effects) was fitted by maximum likelihood (BFGS optimiser) "
    "across the complete 7-division x 38-semester panel (266 division-semester "
    "observations, including explicit zero-outbreak semesters). Poultry density, "
    "being a single static value per division, is perfectly collinear with the "
    "division fixed effects and was therefore not included in this model; its "
    "association with outbreak burden was instead assessed only as a raw "
    "between-division Pearson correlation. A finer-resolution epidemic curve was "
    "reconstructed for the largest tracked WAHIS event (event ID 192, spanning "
    "2007-2013) from the reporting dates and cumulative outbreak counts of its 46 "
    "individual follow-up reports; because this timestamp reflects when a report "
    "was submitted rather than a verified outbreak onset date, this analysis is "
    "presented as report-level, not outbreak-level, temporal resolution. All "
    "analyses were performed in Python 3.12 using pandas, geopandas, rasterio/"
    "rasterstats, and statsmodels 0.14.6."
)

heading("2.5 Ethical Considerations", 2)
para_indent(
    "This study used only publicly available, aggregated animal-health "
    "surveillance data. No individually identifiable human data were accessed at "
    "any stage. Ethical review board approval was not required, consistent with "
    "standard practice for analyses of publicly available, de-identified "
    "veterinary surveillance data."
)

# ══════════════════════════════════════════════════════════════════════════
# 3. RESULTS
# ══════════════════════════════════════════════════════════════════════════

heading("3. Results", 1)

heading("3.1 National Reporting Status, 2007-2025", 2)
para_indent(
    "Bangladesh's six-month country reports show HPAI present in poultry for the "
    "majority of semesters from 2007 through 2013, a return to absence through "
    "2015-2016, and intermittent presence in 2017-2018 and again in 2025 (Figure "
    "1). No six-month report rows exist for 2021 or 2022 -- distinct from an "
    "explicit 'no information' status -- so this period cannot be classified as "
    "either confirmed absence or unreported presence from these data alone."
)

add_figure(FIG_DIR / "fig1_semester_status_timeline.png",
           "Figure 1. HPAI reporting status in Bangladesh by six-month semester, "
           "2007-2025 (WAHIS country reports; poultry and non-poultry/wild-bird "
           "listings shown separately). No rows exist for 2021-2022 (blank, not "
           "'no information'). Source: WAHIS.")

heading("3.2 Outbreak Burden by Division", 2)
para_indent(
    "A total of 564 new HPAI outbreaks were reported across Bangladesh's "
    "administrative divisions from 2007 to 2025 (Table 1, Figures 2 and 3). Dhaka "
    "division accounted for the largest share (266 outbreaks, 47.2% of the "
    "national total; 267 including the single 2025 event resolved to Narayanganj "
    "Sadar), followed by Rajshahi (131, 23.2%) and Chittagong (81, 14.4%). Barisal "
    "(31), Khulna (40), Rangpur (10), and Sylhet (4) each recorded substantially "
    "lower burdens. Across all divisions, WAHIS records show 465,003 susceptible-"
    "animal cases, 2,391,300 birds killed and disposed of under control measures, "
    "and 447,872 recorded deaths."
)

add_table_from_data(
    headers=["Division", "Outbreaks", "Susceptible", "Cases", "Killed/Disposed", "Deaths", "Years Active"],
    rows=[
        ["Dhaka", "266", "1,883,057", "288,286", "1,603,935", "273,406", "2007-2018"],
        ["Rajshahi", "131", "430,180", "52,700", "360,129", "52,235", "2007-2017"],
        ["Chittagong", "81", "294,010", "74,893", "220,708", "73,302", "2007-2012"],
        ["Khulna", "40", "125,456", "27,976", "98,122", "27,781", "2007-2025"],
        ["Barisal", "31", "82,446", "11,817", "70,629", "11,817", "2008-2011"],
        ["Rangpur", "10", "38,848", "7,021", "31,827", "7,021", "2011-2013"],
        ["Sylhet", "4", "8,258", "2,308", "5,950", "2,308", "2008-2011"],
        ["Narayanganj Sadar", "1", "2", "2", "-", "2", "2025"],
    ],
    caption="Table 1. HPAI outbreak burden by administrative division, Bangladesh, "
            "2007-2025. Source: WAHIS event-level quantitative-data export. "
            "'Cases' and 'Killed/Disposed' totals combine poultry and wild-species "
            "records (see Section 3.4)."
)

add_figure(FIG_DIR / "fig2_division_outbreak_burden.png",
           "Figure 2. Total reported HPAI outbreaks by administrative division, "
           "Bangladesh, 2007-2025. Source: WAHIS event-level quantitative data.")

add_figure(FIG_DIR / "fig6_division_choropleth_map.png",
           "Figure 3. Choropleth map of total reported HPAI outbreaks by "
           "administrative division, Bangladesh, 2007-2025. Boundary: geoBoundaries "
           "ADM1, dissolved to WAHIS's 7-division scheme (Mymensingh merged into "
           "Dhaka). Source: WAHIS; geoBoundaries.",
           width=Inches(4.5))

heading("3.3 Poultry Density Does Not Track Outbreak Burden", 2)
para_indent(
    "Zonal-averaged FAO GLW4 (2020) chicken density ranged from 1,217 head/km2 in "
    "Sylhet to 3,394 head/km2 in Barisal -- notably, Barisal has the highest "
    "modelled poultry density of any division yet the fifth-highest outbreak "
    "count (31), while Dhaka, with a mid-range density of 2,230 head/km2, carries "
    "by far the highest outbreak burden. The raw between-division correlation "
    "between outbreak counts and chicken density was negligible (Pearson r = "
    "0.02; Table 2), consistent with reported outbreaks being driven by factors "
    "other than static regional poultry density alone -- plausibly including "
    "live-bird-market and trade-network structure, which this analysis did not "
    "have data to test directly."
)

heading("3.4 Species Involved and the 2025 Wildlife Detection", 2)
para_indent(
    "The great majority of reported cases involved domestic poultry (species "
    "'Birds'). Wild-species detections were first recorded in the WAHIS "
    "quantitative-data export in 2016: House Crow (Corvus splendens) in Rajshahi, "
    "Dhaka, and Khulna divisions across five reporting semesters between 2016 and "
    "2019 (40 to 166 cases per semester), and unidentified Phasianidae in Dhaka in "
    "early 2018 (600 cases) (Figure 4). The most recent WAHIS event for "
    "Bangladesh (event ID 6453, confirmed 2025/04/21) reports H5N1 in a captive "
    "Serval (Leptailurus serval) at Siddirgonj Thana, Narayanganj Sadar, Dhaka "
    "division (23.64696N, 90.512087E; approximate location), with the reason for "
    "notification recorded as 'unusual host species.' Both susceptible animals "
    "were confirmed as cases and both died. The animal category was recorded as "
    "'WILD Captive', indicating a captive-wildlife setting rather than a "
    "free-ranging detection."
)

add_figure(FIG_DIR / "fig3_species_spillover_timeline.png",
           "Figure 4. HPAI cases in Bangladesh by species and semester, 2007-2025. "
           "Top panel: domestic poultry cases (national total per semester). Bottom "
           "panel: wild-species detections (House Crow, Phasianidae, and the 2025 "
           "captive Serval), marker size proportional to reported cases. "
           "Source: WAHIS.")

heading("3.5 A Finer-Resolution View of the 2007-2013 Epidemic Wave", 2)
para_indent(
    "WAHIS tracks Bangladesh's 2007-2013 HPAI poultry epidemic as a single "
    "event (ID 192) followed by 46 individual reports between 30 March 2007 and "
    "21 December 2013, cumulatively reaching 549 outbreaks -- the great majority "
    "of the 2007-2013 national total. Reconstructing the epidemic curve at this "
    "report-level resolution, rather than the coarser six-month semester bins used "
    "elsewhere in this analysis, reveals two distinct peaks rather than one "
    "continuous wave (Figure 5): a sharp spike of 156 new outbreaks reported in "
    "April 2008, and a second, smaller peak of 67 new outbreaks in March 2011, "
    "separated by a comparatively quiet 2009-2010 period. This structure is not "
    "visible in the semester-aggregated data used for the division- and "
    "seasonal-level analyses (Sections 3.2, 3.6)."
)

add_figure(FIG_DIR / "fig4_event192_epidemic_curve.png",
           "Figure 5. Report-level epidemic curve of the 2007-2013 HPAI poultry "
           "wave (WAHIS event 192), reconstructed from 46 individual follow-up "
           "reports. Top: new outbreaks per report. Bottom: cumulative outbreaks. "
           "Two peaks are visible (April 2008, March 2011) that are not "
           "distinguishable in semester-aggregated data. Source: WAHIS.")

heading("3.6 Seasonality and Climate Association", 2)
para_indent(
    "Outbreak occurrence was strongly concentrated in the January-June semester: "
    "525 of 563 division-resolved outbreaks (93.3%) fell in this dry/cool period, "
    "versus 38 (6.7%) in July-December (Table 2b). Mean relative humidity was "
    "correspondingly lower in outbreak-heavy semesters (66.1% in January-June "
    "versus 84.1% in July-December) and mean precipitation substantially lower "
    "(986 mm versus 1,626 mm); mean temperature differed only marginally between "
    "the two halves of the year (25.6 degC versus 25.3 degC) despite the "
    "regression finding a significant temperature association (Section 3.7), "
    "indicating the temperature coefficient should be interpreted as part of a "
    "correlated climate signal rather than in isolation (Section 4.5). Raw "
    "Pearson correlations with semester-level outbreak counts were r = -0.21 for "
    "humidity and r = -0.19 for precipitation (Table 2, Figure 6)."
)

add_table_from_data(
    headers=["Covariate", "Pearson r with new outbreaks"],
    rows=[
        ["Temperature (degC)", "0.070"],
        ["Relative humidity (%)", "-0.213"],
        ["Precipitation (mm)", "-0.193"],
        ["Chicken density (head/km2)", "0.024"],
    ],
    caption="Table 2. Pairwise Pearson correlation between division-semester HPAI "
            "outbreak counts and climate/poultry covariates, 2007-2025 (n = 266 "
            "division-semester observations)."
)

add_figure(FIG_DIR / "fig5_climate_seasonality.png",
           "Figure 6. HPAI outbreaks by division-semester against mean relative "
           "humidity, coloured by dry/cool (Jan-Jun) versus monsoon/warm (Jul-Dec) "
           "semester. Outbreaks concentrate below approximately 75% mean relative "
           "humidity. Source: WAHIS outbreak counts; NASA POWER climate.")

heading("3.7 Negative-Binomial Regression", 2)
para_indent(
    "In a negative-binomial regression with division fixed effects (pseudo R2 = "
    "0.136; likelihood-ratio p = 9.2x10-13; converged), temperature (coefficient "
    "-1.92, p<0.001), relative humidity (-0.14, p = 0.001), and precipitation "
    "(-0.16 per 100 mm, p = 0.016) were each significantly negatively associated "
    "with outbreak counts, confirming the seasonal pattern observed descriptively "
    "in Section 3.6 after controlling for division. Dhaka's division fixed effect "
    "was significantly positive relative to the Barisal baseline (p = 0.042), and "
    "Rangpur and Sylhet were significantly negative (p<0.001 and p = 0.005 "
    "respectively), consistent with the raw division ranking in Table 1. The "
    "three climate covariates were moderately intercorrelated (|r| up to 0.59), "
    "so individual coefficient magnitudes should be interpreted with caution; the "
    "joint significance of the climate block is the more robust finding (Section "
    "4.5)."
)

add_table_from_data(
    headers=["Term", "Coefficient", "IRR", "Std. Err.", "p-value"],
    rows=[
        ["Intercept", "60.52", "-", "16.27", "<0.001"],
        ["Division: Chittagong (ref. Barisal)", "1.07", "2.92", "0.83", "0.195"],
        ["Division: Dhaka", "1.58", "4.85", "0.78", "0.042"],
        ["Division: Khulna", "-0.03", "0.97", "0.74", "0.967"],
        ["Division: Rajshahi", "0.13", "1.14", "0.71", "0.855"],
        ["Division: Rangpur", "-4.54", "0.011", "1.25", "<0.001"],
        ["Division: Sylhet", "-3.48", "0.031", "1.25", "0.005"],
        ["Temperature (degC)", "-1.92", "0.147", "0.55", "<0.001"],
        ["Humidity (%)", "-0.14", "0.872", "0.04", "0.001"],
        ["Precipitation (per 100mm)", "-0.16", "0.848", "0.07", "0.016"],
        ["alpha (dispersion)", "6.37", "-", "1.29", "<0.001"],
    ],
    caption="Table 3. Negative-binomial regression of HPAI outbreak counts on "
            "climate covariates with division fixed effects (n = 266 division-"
            "semester observations; reference division = Barisal). IRR = "
            "incidence rate ratio, exp(coefficient)."
)

# ══════════════════════════════════════════════════════════════════════════
# 4. DISCUSSION
# ══════════════════════════════════════════════════════════════════════════

heading("4. Discussion", 1)

heading("4.1 A Persistent, Dhaka-Centred Poultry-Sector Disease", 2)
para_indent(
    "Two decades of WAHIS surveillance confirm HPAI as a recurring, rather than "
    "one-off, feature of Bangladesh's poultry sector, with Dhaka division "
    "consistently the epicentre. Dhaka's dominance (47.2% of the national "
    "outbreak total, and a significantly elevated division fixed effect even "
    "after accounting for climate) most plausibly reflects its position as "
    "Bangladesh's largest live-bird-market and poultry-trade hub rather than "
    "simply its poultry density, which was mid-range rather than highest among "
    "divisions (Section 3.3). This is consistent with the well-established role "
    "of live-bird markets, rather than raw farm density alone, as amplification "
    "points for HPAI transmission in South and Southeast Asia [1,3]."
)

heading("4.2 Strong Dry-Season Seasonality", 2)
para_indent(
    "The concentration of outbreaks in the dry, cooler January-June semester "
    "(93.3% of division-resolved outbreaks) and its confirmation in a regression "
    "controlling for division is consistent with the broader literature on "
    "avian-influenza-virus persistence, which favours cooler, lower-humidity "
    "conditions for environmental survival [2], and with the timing of "
    "Bangladesh's winter migratory wild-bird season and heightened live-poultry "
    "trade around winter festivals. This seasonal predictability, if confirmed in "
    "future years, could support timing enhanced surveillance and biosecurity "
    "messaging ahead of the dry season rather than distributing effort evenly "
    "across the year."
)

heading("4.3 Two Peaks Within One Tracked Epidemic", 2)
para_indent(
    "The report-level reconstruction of WAHIS event 192 (Section 3.5) illustrates "
    "how semester-aggregated surveillance data can obscure epidemic structure: "
    "what appears as one continuous 2007-2013 wave in six-month reports resolves, "
    "at near-monthly report resolution, into two distinct peaks (April 2008 and "
    "March 2011) separated by a comparatively quiet two-year period. This "
    "distinction matters for interpreting whether control measures, re-"
    "introduction, or genuinely independent epidemic drivers explain the 2011 "
    "resurgence -- a question this descriptive analysis cannot resolve but that "
    "the underlying WAHIS report history (data/processed/hpai_event192_report_"
    "history.csv in the accompanying repository) could support with further "
    "investigation."
)

heading("4.4 The 2025 Captive-Serval Detection: A One Health Signal", 2)
para_indent(
    "The most recent WAHIS record for Bangladesh -- H5N1 in a captive Serval, "
    "with both susceptible animals dying -- is notable less for its scale (a "
    "single small event) than for what it represents: a confirmed detection in a "
    "non-poultry, non-avian host, explicitly flagged by the reporting authority "
    "as an 'unusual host species' event. Internationally, H5N1 clade 2.3.4.4b has "
    "caused widening mammalian spillover since 2020, including in captive and "
    "wild felids in multiple countries [8]; this analysis cannot determine "
    "whether the Bangladesh Serval case belongs to that clade, as WAHIS did not "
    "report a full genotype for this event, but the event's classification as "
    "'WILD Captive' -- most consistent with a zoo, rescue, or private captive-"
    "wildlife setting -- points to captive-wildlife biosecurity, rather than "
    "confirmed free-ranging wildlife transmission, as the immediate area for "
    "follow-up. Combined with the 2016-2019 House Crow and Phasianidae "
    "detections, this case underscores the value of maintaining wild/non-poultry "
    "HPAI surveillance in Bangladesh even during years when poultry-sector "
    "reporting is quiet."
)

heading("4.5 Limitations", 2)
para_indent(
    "Several limitations should be considered. First, WAHIS-reported outbreaks "
    "reflect official government reporting, which is known globally to "
    "underrepresent true incidence, particularly in informal backyard and "
    "smallholder poultry sectors; the absence of report rows for 2021-2022 in "
    "particular cannot be distinguished from a true reporting gap versus genuine "
    "disease absence. Second, spatial resolution in this analysis is limited to "
    "administrative division: WAHIS's individual-outbreak records do carry exact "
    "coordinates (confirmed for the 2025 Serval case, Section 3.4), but no bulk "
    "export of this coordinate layer exists, and the largest event alone (192) "
    "comprises 549 individual outbreaks, making point-level extraction "
    "impractical without direct WAHIS data-sharing access; point-pattern or "
    "kernel-density analysis was therefore out of scope. Third, poultry density "
    "is a single 2020 cross-section applied uniformly across 2007-2025 and does "
    "not capture the sector's substantial growth over this period, nor "
    "distinguish commercial from backyard production. Fourth, division-centroid "
    "climate values are coarse approximations of conditions at actual outbreak "
    "sites. Fifth, the three climate covariates are moderately intercorrelated "
    "(Section 3.7), so the regression should be read as confirming a joint "
    "seasonal climate signal rather than precisely apportioning independent "
    "effects to temperature, humidity, and precipitation individually. Sixth, "
    "this is an ecological, ill-designed-for-causality analysis at the division-"
    "semester level; individual-level or farm-level risk factors (bird density on "
    "individual holdings, market connectivity, biosecurity practices) that prior "
    "Bangladesh-specific studies have examined for the 2007-2013 period [3,4] "
    "were outside the scope of this national surveillance-record analysis."
)

# ══════════════════════════════════════════════════════════════════════════
# 5. CONCLUSIONS
# ══════════════════════════════════════════════════════════════════════════

heading("5. Conclusions", 1)

para_indent(
    "Highly pathogenic avian influenza has recurred in Bangladesh's poultry "
    "sector across nearly two decades of WAHIS surveillance, concentrated in "
    "Dhaka division and strongly seasonal toward the dry, cooler months -- a "
    "pattern that persists after controlling for division and is not explained "
    "by poultry density alone. Report-level reconstruction of the largest tracked "
    "epidemic wave reveals internal structure (two peaks, not one) invisible in "
    "standard semester-aggregated reporting. The 2025 detection of H5N1 in a "
    "captive Serval, alongside earlier wild-bird detections, signals an emerging "
    "One Health surveillance need in Bangladesh that extends beyond the "
    "traditional poultry-sector focus of prior national HPAI research."
)

# ══════════════════════════════════════════════════════════════════════════
# DATA AVAILABILITY
# ══════════════════════════════════════════════════════════════════════════

heading("Data Availability Statement", 1)

para(
    "All data, cleaning/analysis code, and figures used in this study are "
    "publicly available at: "
    "https://github.com/khalilurrrahmanridoykhan/bangladesh-avian-influenza-2007-2025. "
    "Source data were retrieved from the WOAH World Animal Health Information "
    "System (https://wahis.woah.org), the NASA POWER API "
    "(https://power.larc.nasa.gov), and the FAO Gridded Livestock of the World "
    "v4 catalog (https://data.apps.fao.org)."
)

heading("Competing Interests", 1)
para("The author declares no competing interests.")

heading("Funding", 1)
para("This research received no specific funding from any funding agency in the "
     "public, commercial, or not-for-profit sectors.")

# ══════════════════════════════════════════════════════════════════════════
# REFERENCES
# ══════════════════════════════════════════════════════════════════════════

page_break()
heading("References", 1)

para(
    "NOTE TO AUTHOR: references [3], [4], and [8] below are cited from memory "
    "and MUST be independently verified (author names, year, volume/page, DOI) "
    "against the original sources before submission. References [1], [2], [5], "
    "[6], and [7] point to organizational/data sources used directly by this "
    "analysis pipeline and are safe to cite as-is.",
    italic=True, space_after=12,
)

refs = [
    "[1] WOAH (World Organisation for Animal Health). Avian Influenza. "
    "Disease fact sheet. Available at: https://www.woah.org/en/disease/"
    "avian-influenza/",

    "[2] FAO (Food and Agriculture Organization of the United Nations). "
    "H5N1 highly pathogenic avian influenza global situation update. "
    "FAO Emergency Prevention System (EMPRES) for animal health. Available at: "
    "https://www.fao.org/animal-health/situation-updates/en",

    "[3] Biswas PK, Christensen JP, Ahmed SS, et al. Avian influenza outbreaks "
    "in chickens, Bangladesh. Emerging Infectious Diseases. 2008;14(12):"
    "1909-1912. [VERIFY BEFORE SUBMISSION]",

    "[4] Loth L, Gilbert M, Osmani MG, Kalam AM, Xiao X. Risk factors and "
    "clusters of highly pathogenic avian influenza H5N1 outbreaks in "
    "Bangladesh. Preventive Veterinary Medicine. 2010;96(1-2):104-113. "
    "[VERIFY BEFORE SUBMISSION]",

    "[5] geoBoundaries. Bangladesh ADM1 administrative boundaries "
    "(2015 vintage), CC0 1.0. Runfola D, et al. geoBoundaries: A global "
    "database of political administrative boundaries. PLOS ONE. "
    "2020;15(4):e0231866. Available at: https://www.geoboundaries.org",

    "[6] NASA Langley Research Center POWER Project. NASA POWER (Prediction "
    "of Worldwide Energy Resources) daily point API. Available at: "
    "https://power.larc.nasa.gov",

    "[7] Gilbert M, Nicolas G, Cinardi G, et al. Global distribution data "
    "for cattle, buffaloes, horses, sheep, goats, pigs, chickens and ducks "
    "in 2010 (Gridded Livestock of the World). Scientific Data. "
    "2018;5:180227. GLW4 (2020, 10km) accessed via FAO catalog: "
    "https://data.apps.fao.org [VERIFY GLW4-specific citation before submission]",

    "[8] Mostafa A, Naguib MM, Nogales A, et al. Avian influenza A(H5N1) "
    "virus clade 2.3.4.4b spillover into mammals. [VERIFY EXACT SOURCE, "
    "AUTHORS, YEAR BEFORE SUBMISSION -- cited from memory as general "
    "background on mammalian spillover of clade 2.3.4.4b.]",

    "[9] Bangladesh Bureau of Statistics (BBS). Population and Housing "
    "Census 2022: National Report. Statistics and Informatics Division, "
    "Ministry of Planning, Bangladesh; 2023.",

    "[10] WOAH. World Animal Health Information System (WAHIS). Available "
    "at: https://wahis.woah.org. Accessed August 2026.",
]

for ref in refs:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.first_line_indent = Cm(-1.0)
    p.paragraph_format.line_spacing = Pt(18)
    run = p.add_run(ref)
    run.font.name = "Times New Roman"
    run.font.size = Pt(11)

# Figures are now placed inline in the Results section, immediately after
# the paragraph that discusses each one, rather than batched at the end.

# ── Save ─────────────────────────────────────────────────────────────────
doc.save(OUT)
print(f"\nManuscript saved: {OUT}")
print(f"File size: {OUT.stat().st_size / 1024:.0f} KB")
print("\n*** IMPORTANT: References [3], [4], and [8] are cited from memory and")
print("*** MUST be independently verified before this manuscript is submitted anywhere.")
